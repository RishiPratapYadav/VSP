import json
import io
from pathlib import Path
from docx import Document
from PyPDF2 import PdfReader
from fastapi import UploadFile

# --- Files / folders ---
GLOBAL_COUNTER_FILE = Path("global/global_counter.json")
SUBMISSION_FOLDER = Path("data/submissions")
RFP_FOLDER = Path("data/rfps")
SCHEMA_DIR = Path("schema")
VENDOR_FOLDER = Path("data/vendor_responses")

# --- Create directories ---
for folder in [SUBMISSION_FOLDER, RFP_FOLDER, VENDOR_FOLDER, GLOBAL_COUNTER_FILE.parent, SCHEMA_DIR]:
    folder.mkdir(parents=True, exist_ok=True)


def get_next_initiative_id() -> int:
    if not GLOBAL_COUNTER_FILE.exists():
        counter = {"last_id": 0}
    else:
        with open(GLOBAL_COUNTER_FILE, "r") as f:
            try:
                counter = json.load(f)
            except json.JSONDecodeError:
                counter = {"last_id": 0}
    counter["last_id"] = int(counter.get("last_id", 0)) + 1
    with open(GLOBAL_COUNTER_FILE, "w") as f:
        json.dump(counter, f, indent=2)
    return counter["last_id"]


def save_submission(initiative_id: int, data: dict, schema_name: str = None):
    if schema_name:
        file_path = SUBMISSION_FOLDER / f"initiative_{initiative_id}_{schema_name}.json"
    else:
        file_path = SUBMISSION_FOLDER / f"initiative_{initiative_id}.json"
    with open(file_path, "w") as f:
        json.dump(data, f, indent=2)


def load_initiative_data(initiative_id: int, schema_name: str) -> dict:
    """Loads and merges the base and detailed submission data for an initiative."""
    base_file = SUBMISSION_FOLDER / f"initiative_{initiative_id}.json"
    details_file = SUBMISSION_FOLDER / f"initiative_{initiative_id}_{schema_name}.json"

    if not base_file.exists() or not details_file.exists():
        raise FileNotFoundError("Initiative data files not found.")

    with open(base_file, "r") as f:
        base_data = json.load(f)
    with open(details_file, "r") as f:
        details_data = json.load(f)

    return {**base_data, **details_data}


def save_rfp_doc(text: str, initiative_id: int) -> str:
    output_file = RFP_FOLDER / f"initiative_{initiative_id}_rfp.docx"
    doc = Document()
    for line in text.splitlines():
        l = line.strip()
        if l.startswith("###"):
            doc.add_heading(l.lstrip("# ").strip(), level=3)
        elif l.startswith("##"):
            doc.add_heading(l.lstrip("# ").strip(), level=2)
        elif l.startswith("#"):
            doc.add_heading(l.lstrip("# ").strip(), level=1)
        elif not l:
            doc.add_paragraph("")
        else:
            doc.add_paragraph(l)
    doc.save(output_file)
    return str(output_file)


async def save_vendor_files(initiative_id: int, files: list[UploadFile]):
    upload_dir = VENDOR_FOLDER / f"initiative_{initiative_id}"
    upload_dir.mkdir(parents=True, exist_ok=True)

    combined_data = {}
    for file in files:
        content = await file.read()
        text = ""
        if file.filename.lower().endswith(".pdf"):
            pdf = PdfReader(io.BytesIO(content))
            text = "".join(page.extract_text() or "" for page in pdf.pages)
        elif file.filename.lower().endswith(".docx"):
            doc = Document(io.BytesIO(content))
            text = "\n".join(para.text for para in doc.paragraphs)
        else:
            text = content.decode("utf-8", errors="ignore")

        combined_data[file.filename] = text.strip()
        with open(upload_dir / file.filename, "wb") as f:
            f.write(content)

    combined_path = upload_dir / "combined_vendor_responses.json"
    with open(combined_path, "w") as f:
        json.dump(combined_data, f, indent=2)