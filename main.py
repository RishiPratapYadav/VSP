# app.py
import os
import json
from datetime import date
from pathlib import Path
from typing import Dict, Any
from fastapi import FastAPI, Request, UploadFile, File
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, RedirectResponse
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
import html as html_lib
import io
import google.generativeai as genai
from PyPDF2 import PdfReader

# --- Config ---
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    print("WARNING: GOOGLE_API_KEY environment variable not set. AI features will not work.")
    gemini_model = None
else:
    genai.configure(api_key=GOOGLE_API_KEY)
    gemini_model = genai.GenerativeModel('gemini-2.0-pro-exp')

app = FastAPI()

# --- Files / folders ---
GLOBAL_COUNTER_FILE = Path("global/global_counter.json")
SUBMISSION_FOLDER = Path("data/submissions")
RFP_FOLDER = Path("data/rfps")
SCHEMA_FILE = Path("schema/form_schema.json")
RFP_TEMPLATE_FOLDER = Path("templates/rfp_templates")
SCHEMA_DIR = Path("schema")

SUBMISSION_FOLDER.mkdir(parents=True, exist_ok=True)
RFP_FOLDER.mkdir(parents=True, exist_ok=True)
GLOBAL_COUNTER_FILE.parent.mkdir(parents=True, exist_ok=True)
RFP_TEMPLATE_FOLDER.mkdir(parents=True, exist_ok=True)
SCHEMA_DIR.mkdir(parents=True, exist_ok=True)

# Example schema map - update with your schema filenames
SCHEMA_MAP = {
    ("Clinical", "Manufacturing"): "clinical_manufacturing.json",
    ("Clinical", "Testing"): "clinical_testing.json",
    ("Clinical", "Packaging"): "clinical_packaging.json",
    ("Commercial", "Manufacturing"): "commercial_manufacturing.json",
    ("Commercial", "Packaging"): "commercial_packaging.json"
}

# --- Styling and helpers for UI ---
STYLE = """
<style>
body { font-family: 'Segoe UI', Tahoma, sans-serif; background:#f6f8fa; color:#222; margin:0; padding:0; display: flex; }
.sidebar { width: 240px; background: #2c3e50; color: #ecf0f1; padding: 20px; height: 100vh; position: fixed; }
.sidebar h2 { color: #ecf0f1; border: none; }
.sidebar a { color: #ecf0f1; text-decoration: none; display: block; padding: 10px 15px; border-radius: 4px; margin-bottom: 8px; }
.sidebar a:hover, .sidebar a.active { background-color: #34495e; }
.main-content { margin-left: 280px; padding: 20px; width: calc(100% - 280px); }
.container { max-width:1000px; margin:12px auto; background:#fff; padding:28px; border-radius:12px; box-shadow:0 6px 24px rgba(0,0,0,0.08); }
h1 { color:#1a73e8; margin:0 0 10px 0; text-align:center;}
h2 { color:#333; margin-top:18px; border-bottom:1px solid #eee; padding-bottom:8px }
.form-grid { display:grid; grid-template-columns:1fr 1fr; gap:16px; }
@media(max-width:800px){ .form-grid{grid-template-columns:1fr} }
label { display:block; font-weight:600; margin-top:10px; }
input, select, textarea { width:100%; padding:10px; margin-top:6px; border-radius:8px; border:1px solid #d1d7e0; font-size:15px; }
textarea { min-height:110px; resize: vertical; }
.checkbox-group,label.radio-group { margin-top:8px; }
button { background:#1a73e8; color:#fff; padding:12px 18px; border-radius:8px; border:none; cursor:pointer; margin-top:18px; font-weight:700; }
button:hover { background: #155ab3; }
.progress { display:flex; gap:12px; margin-bottom:18px; justify-content:space-between; }
.step { flex:1; text-align:center; color:#a2abb8; font-weight:700; position:relative; padding:8px 6px; }
.step.active { color:#1a73e8; }
.step::before { content:'●'; display:block; font-size:20px; margin-bottom:6px; }
.step:not(:last-child)::after { content:''; position:absolute; right:-6px; top:22px; width:12px; height:4px; background:#e6e9ee; z-index:-1; }
.step.active:not(:last-child)::after { background:#1a73e8; }
.loader { border:4px solid #f3f3f3; border-top:4px solid #1a73e8; border-radius:50%; width:44px; height:44px; animation:spin 1s linear infinite; margin:28px auto; }
@keyframes spin { 0%{transform:rotate(0)}100%{transform:rotate(360deg)} }
.rfp-output { background:#fafbfd; padding:18px; border-radius:10px; white-space:pre-wrap; line-height:1.5; font-family:ui-monospace, SFMono-Regular, Menlo, Monaco, Monaco, "Roboto Mono", "Courier New", monospace; }
.download { display:inline-block; margin-top:18px; background:#22a66b; color:#fff; padding:10px 16px; border-radius:8px; text-decoration:none; }
.download:hover { background: #20945d; }
.notice { margin-top:12px; color:#666; font-size:14px }
.initiative-list { list-style-type: none; padding: 0; }
.initiative-list li { background: #fdfdfd; border: 1px solid #eee; padding: 15px; margin-bottom: 10px; border-radius: 8px; display: flex; justify-content: space-between; align-items: center; }
.initiative-list .info { font-size: 16px; }
.initiative-list .info strong { color: #1a73e8; }
.initiative-list .actions a { margin-left: 10px; font-size: 14px; }
</style>
"""

def get_base_layout(title: str, content: str) -> HTMLResponse:
    html = f"""<!DOCTYPE html><html><head><title>{html_lib.escape(title)}</title>{STYLE}</head><body>
    <div class="sidebar"><h2>RFP Assistant</h2><nav><a href="/">New Vendor Request</a><a href="/initiatives">List Initiatives</a></nav></div>
    <main class="main-content">{content}</main></body></html>"""
    return HTMLResponse(content=html)

def render_progress(step:int):
    steps = ["1. Basic Info", "2. Details & Scoring", "3. Generate RFP"]
    html = '<div class="progress">'
    for i, label in enumerate(steps, start=1):
        cls = "step active" if i <= step else "step"
        html += f'<div class="{cls}">{label}</div>'
    html += "</div>"
    return html

# --- Counter helper ---
def get_next_initiative_id() -> int:
    if not GLOBAL_COUNTER_FILE.exists():
        counter = {"last_id": 0}
    else:
        with open(GLOBAL_COUNTER_FILE, "r") as f:
            try:
                counter = json.load(f)
            except Exception:
                counter = {"last_id": 0}
    counter["last_id"] = int(counter.get("last_id", 0)) + 1
    with open(GLOBAL_COUNTER_FILE, "w") as f:
        json.dump(counter, f, indent=2)
    return counter["last_id"]

# --- Schema loader & HTML form generator (flat "fields" with "section") ---
def load_schema(file_name: str):
    path = SCHEMA_DIR / file_name
    if path.exists():
        with open(path, "r") as f:
            return json.load(f)
    return None

def generate_form_html(schema: Dict[str,Any], action="/submit", defaults:Dict[str,Any]=None):
    defaults = defaults or {}
    # group fields by section
    sections: Dict[str, list] = {}
    for field in schema.get("fields", []):
        section = field.get("section","General")
        sections.setdefault(section, []).append(field)

    html = '<div class="container">'
    html += render_progress(1)
    html += f"<h1>{html_lib.escape(schema.get('title', 'Form'))}</h1>"
    html += f'<form method="post" action="{action}">'

    # Render sections and fields
    for section, fields in sections.items():
        html += f"<h2>{html_lib.escape(section)}</h2>"
        # simple grid for two-column layout
        html += '<div class="form-grid">'
        for field in fields:
            name = field.get("name")
            label = field.get("label", name)
            ftype = field.get("type", "text")
            default = defaults.get(name, field.get("default", ""))

            html += "<div>"
            html += f'<label for="{html_lib.escape(name)}">{html_lib.escape(label)}</label>'

            if ftype in ("text","email","tel","number"):
                val = html_lib.escape(str(default)) if default is not None else ""
                required = "required" if field.get("required", False) else ""
                html += f'<input type="{ftype}" name="{html_lib.escape(name)}" value="{val}" {required}>'
            elif ftype == "textarea":
                val = html_lib.escape(str(default)) if default is not None else ""
                html += f'<textarea name="{html_lib.escape(name)}">{val}</textarea>'
            elif ftype == "select":
                html += f'<select name="{html_lib.escape(name)}">'
                for opt in field.get("options",[]):
                    sel = 'selected' if str(opt) == str(default) else ''
                    html += f'<option value="{html_lib.escape(str(opt))}" {sel}>{html_lib.escape(str(opt))}</option>'
                html += '</select>'
            elif ftype == "checkbox":
                # multiple checkboxes (name repeated)
                html += '<div class="checkbox-group">'
                for opt in field.get("options",[]):
                    checked = ""
                    # default could be list or comma separated string
                    if isinstance(default, list) and str(opt) in [str(x) for x in default]:
                        checked = "checked"
                    elif isinstance(default, str) and str(opt) in default.split(","):
                        checked = "checked"
                    html += f'<label><input type="checkbox" name="{html_lib.escape(name)}" value="{html_lib.escape(str(opt))}" {checked}> {html_lib.escape(str(opt))}</label>'
                html += '</div>'
            elif ftype == "radio":
                html += '<div class="radio-group">'
                for opt in field.get("options",[]):
                    checked = "checked" if str(opt) == str(default) else ""
                    html += f'<label><input type="radio" name="{html_lib.escape(name)}" value="{html_lib.escape(str(opt))}" {checked}> {html_lib.escape(str(opt))}</label>'
                html += '</div>'
            else:
                html += f'<input type="text" name="{html_lib.escape(name)}" value="{html_lib.escape(str(default))}">'
            html += "</div>"  # field column
        html += '</div>'  # grid
    html += '<button type="submit">Continue</button></form></div>'
    return html

# --- Save DOCX helper ---
def save_rfp_doc(text: str, initiative_id: int) -> str:
    RFP_FOLDER.mkdir(parents=True, exist_ok=True)
    output_file = RFP_FOLDER / f"initiative_{initiative_id}_rfp.docx"
    doc = Document()
    # Simple markup: lines starting with ### or ## or # as headings
    for line in text.splitlines():
        l = line.strip()
        if l.startswith("###"):
            doc.add_heading(l.lstrip("# ").strip(), level=3)
        elif l.startswith("##"):
            doc.add_heading(l.lstrip("# ").strip(), level=2)
        elif l.startswith("#"):
            doc.add_heading(l.lstrip("# ").strip(), level=1)
        elif l == "":
            doc.add_paragraph("")  # blank line
        else:
            doc.add_paragraph(l)
    doc.save(output_file)
    return str(output_file)

def save_comparison_docx(data: dict, initiative_id: int) -> str:
    """Saves the vendor comparison data to a .docx file."""
    output_file = VENDOR_FOLDER / f"initiative_{initiative_id}" / "comparison_result.docx"
    doc = Document()
    doc.add_heading(f"Vendor Comparison for Initiative #{initiative_id}", level=1)

    if "recommendation" in data:
        doc.add_heading("Overall Recommendation", level=2)
        doc.add_paragraph(data["recommendation"].get("summary", "No summary provided."))
        doc.add_paragraph("Top Vendors: " + ", ".join(data["recommendation"].get("top_vendors", ["N/A"])))

    for vendor in data.get("vendors", []):
        doc.add_heading(vendor.get("vendor_name", "Unknown Vendor"), level=2)

        doc.add_heading("Key Proposal Points", level=3)
        doc.add_paragraph(vendor.get("summary", "Not available."))

        doc.add_heading("Evaluation Scores", level=3)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Criterion'
        hdr_cells[1].text = 'Score (/10)'
        hdr_cells[2].text = 'Percentage'

        for criterion, score_details in vendor.get("scores", {}).items():
            row_cells = table.add_row().cells
            row_cells[0].text = criterion
            row_cells[1].text = str(score_details.get("score", "N/A"))
            row_cells[2].text = f"{score_details.get('percentage', 'N/A')}%"

        doc.add_heading("Strengths", level=3)
        doc.add_paragraph(vendor.get("strengths", "Not available."))

        doc.add_heading("Weaknesses", level=3)
        doc.add_paragraph(vendor.get("weaknesses", "Not available."))

        doc.add_heading("Risks / Alignment", level=3)
        doc.add_paragraph(vendor.get("risks", "Not available."))

        doc.add_paragraph() # Add space between vendors

    doc.save(output_file)
    return str(output_file)

def save_comparison_xlsx(data: dict, initiative_id: int) -> str:
    """Saves the vendor comparison data to an .xlsx file."""
    output_file = VENDOR_FOLDER / f"initiative_{initiative_id}" / "comparison_result.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Vendor Comparison"

    headers = ["Vendor Name", "Criterion", "Score (/10)", "Percentage", "Summary", "Strengths", "Weaknesses", "Risks"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')

    for vendor in data.get("vendors", []):
        vendor_name = vendor.get("vendor_name", "Unknown Vendor")
        for criterion, score_details in vendor.get("scores", {}).items():
            ws.append([
                vendor_name,
                criterion,
                score_details.get("score"),
                score_details.get("percentage"),
                vendor.get("summary"),
                vendor.get("strengths"),
                vendor.get("weaknesses"),
                vendor.get("risks"),
            ])
    wb.save(output_file)
    return str(output_file)

# --- Data loading helper ---
def load_initiative_data(initiative_id: int, schema_name: str) -> dict:
    """Loads and merges the base and detailed submission data for an initiative."""
    base_file = SUBMISSION_FOLDER / f"initiative_{initiative_id}.json"
    details_file = SUBMISSION_FOLDER / f"initiative_{initiative_id}_{schema_name}.json"

    if not base_file.exists() or not details_file.exists():
        raise FileNotFoundError("Initiative data files not found.")

    with open(base_file, "r") as f:
        base_data = json.load(f)
    with open(details_file, "r") as f:
        details_data = json.load(f)

    # Merge the two dictionaries
    return {**base_data, **details_data}

@app.get("/", response_class=HTMLResponse)
async def main_form():
    if not SCHEMA_FILE.exists():
        return HTMLResponse("<h3>No main schema found at schema/form_schema.json</h3>", status_code=500)
    schema = load_schema(SCHEMA_FILE.name)
    html = generate_form_html(schema, action="/submit")
    return get_base_layout("New Initiative", html)

@app.get("/edit/{initiative_id}", response_class=HTMLResponse)
async def edit_initiative_form(initiative_id: int):
    """Displays the main form pre-filled with an initiative's data for editing."""
    base_file = SUBMISSION_FOLDER / f"initiative_{initiative_id}.json"
    if not base_file.exists():
        return HTMLResponse("<h3>Initiative not found.</h3>", status_code=404)

    with open(base_file, "r") as f:
        defaults = json.load(f)

    schema = load_schema(SCHEMA_FILE.name)
    if not schema:
        return HTMLResponse("<h3>No main schema found at schema/form_schema.json</h3>", status_code=500)

    action_url = f"/update/{initiative_id}"
    html = generate_form_html(schema, action=action_url, defaults=defaults)
    return get_base_layout(f"Edit Initiative #{initiative_id}", html)

@app.get("/initiatives", response_class=HTMLResponse)
async def list_initiatives():
    """Lists all created initiatives."""
    initiatives = []
    for file in sorted(SUBMISSION_FOLDER.glob("initiative_*.json"), reverse=True):
        if "_" in file.stem.replace("initiative_", ""):
            continue # Skip detailed schema files
        try:
            with open(file, "r") as f:
                data = json.load(f)
                initiatives.append(data)
        except (json.JSONDecodeError, KeyError):
            continue

    list_html = "<h1>📝 All Initiatives</h1>"
    if not initiatives:
        list_html += "<p>No initiatives found. <a href='/'>Create one now</a>.</p>"
    else:
        list_html += '<ul class="initiative-list">'
        for init in initiatives:
            init_id = init.get("initiative_id")
            req_type = init.get("request_type", "N/A")
            services = init.get("services_needed", "N/A")
            if isinstance(services, list):
                services = services[0] if services else "N/A"
            schema_key = (req_type, services)
            schema_file = SCHEMA_MAP.get(schema_key)
            schema_name = schema_file.replace('.json', '') if schema_file else None

            list_html += f'<li><div class="info">Initiative <strong>#{init_id}</strong> &mdash; {req_type} / {services}</div>'
            list_html += '<div class="actions">'
            list_html += f'<a href="/edit/{init_id}">Edit</a>'
            if schema_name:
                list_html += f'<a href="/rfp/{init_id}/{schema_name}">Generate RFP</a>'
            list_html += f'<a href="/upload_vendor_responses/{init_id}">Upload Responses</a>'
            list_html += f'<a href="/compare_vendors/{init_id}">Compare</a>'
            list_html += '</div></li>'
        list_html += '</ul>'

    container_html = f'<div class="container">{list_html}</div>'
    return get_base_layout("All Initiatives", container_html)

@app.post("/submit", response_class=HTMLResponse)
async def submit_main(request: Request):
    form = await request.form()
    # convert MultiDict -> dict with lists for repeated names (checkbox)
    data = {}
    for k, v in form.multi_items():
        if k in data:
            if isinstance(data[k], list):
                data[k].append(v)
            else:
                data[k] = [data[k], v]
        else:
            data[k] = v

    initiative_id = get_next_initiative_id()
    data["initiative_id"] = initiative_id
    base_file = SUBMISSION_FOLDER / f"initiative_{initiative_id}.json"
    with open(base_file, "w") as f:
        json.dump(data, f, indent=2)

    # Decide next schema based on request_type + services_needed
    request_type = data.get("request_type")
    services_needed = data.get("services_needed")
    # For checkbox lists, ensure first if list is present
    if isinstance(services_needed, list):
        services_needed = services_needed[0] if services_needed else None

    schema_file = SCHEMA_MAP.get((request_type, services_needed))
    if not schema_file:
        # no specific schema -> show a simple confirmation with link to view
        html = '<div class="container">'
        html += render_progress(1)
        html += f"<h1>Submission saved — Initiative {initiative_id}</h1>"
        html += '<p class="notice">No detailed schema mapped for this Request Type & Service. You can:</p>'
        html += f'<p><a href="/initiative/{initiative_id}">View submission JSON</a></p>'
        html += "</div>"
        return get_base_layout(f"Initiative #{initiative_id}", html)

    # load schema and render details form
    schema = load_schema(schema_file)
    if not schema:
        return HTMLResponse(f"<h3>Schema file {schema_file} missing in schema/ folder.</h3>", status_code=500)

    # render details form; ensure action posts to /submit/{schema_name}/{initiative_id}
    action = f"/submit/{schema_file.replace('.json','')}/{initiative_id}"
    form_html = generate_form_html(schema, action=action)
    # update progress to step 2 (generate_form_html currently uses render_progress(1) by default).
    # We'll replace the progress with step 2 manually:
    form_html = form_html.replace(render_progress(1), render_progress(2))
    return get_base_layout(f"Details for Initiative #{initiative_id}", form_html)

@app.post("/update/{initiative_id}", response_class=HTMLResponse)
async def update_initiative(request: Request, initiative_id: int):
    """Handles updates for the main initiative form."""
    form = await request.form()
    data = {}
    for k, v in form.multi_items():
        if k in data:
            if isinstance(data[k], list):
                data[k].append(v)
            else:
                data[k] = [data[k], v]
        else:
            data[k] = v

    data["initiative_id"] = initiative_id  # Ensure the ID remains the same
    base_file = SUBMISSION_FOLDER / f"initiative_{initiative_id}.json"
    with open(base_file, "w") as f:
        json.dump(data, f, indent=2)

    return RedirectResponse(url="/initiatives", status_code=303)

@app.post("/submit/{schema_name}/{initiative_id}", response_class=HTMLResponse)
async def submit_details(request: Request, schema_name: str, initiative_id: int):
    form = await request.form()
    data = {}
    for k, v in form.multi_items():
        if k in data:
            if isinstance(data[k], list):
                data[k].append(v)
            else:
                data[k] = [data[k], v]
        else:
            data[k] = v

    # Ensure submissions folder exists
    SUBMISSION_FOLDER.mkdir(parents=True, exist_ok=True)
    file_path = SUBMISSION_FOLDER / f"initiative_{initiative_id}_{schema_name}.json"
    with open(file_path, "w") as f:
        json.dump(data, f, indent=2)

    # Confirm and provide link to generate RFP
    html = '<div class="container">'
    html += render_progress(2)
    html += f"<h1>Details saved for Initiative {initiative_id}</h1>"
    html += f'<p class="notice">You can now generate the RFP enhanced by the AI.</p>'
    html += f'<a class="download" style="background-color:#3367D6;" href="/find_vendors/{initiative_id}/{schema_name}">🔍 Find Vendors</a>'
    html += f'<a class="download" href="/rfp/{initiative_id}/{schema_name}">Generate RFP</a>'
    html += f'<p style="margin-top:10px;"><a href="/initiatives">⟵ Back to All Initiatives</a></p>'
    html += "</div>"
    return get_base_layout(f"Initiative #{initiative_id} Saved", html)

@app.get("/rfp/{initiative_id}/{schema_name}", response_class=HTMLResponse)
async def rfp_loading(initiative_id:int, schema_name:str):
    # show a loading screen, then redirect to result page which actually runs ollama
    html = '<div class="container">'
    html += render_progress(3)
    html += "<h1>⏳ Generating RFP...</h1>"
    html += '<div class="loader"></div>'
    # Short JS redirect to result endpoint after small delay
    html += f"""
    <script>
      setTimeout(function(){{ window.location.href = '/rfp_result/{initiative_id}/{schema_name}'; }}, 800);
    </script>
    """
    html += '</div>'
    return get_base_layout("Generating RFP...", html)

@app.get("/rfp_result/{initiative_id}/{schema_name}", response_class=HTMLResponse)
async def rfp_result(initiative_id:int, schema_name:str):
    try:
        initiative_data = load_initiative_data(initiative_id, schema_name)
    except FileNotFoundError:
        return HTMLResponse("<h3>Initiative files not found. Make sure both JSON submissions exist.</h3>", status_code=404)

    rfp_text = ""
    template_path = RFP_TEMPLATE_FOLDER / f"{schema_name}.txt"
    source_notice = ""

    if template_path.exists():
        # --- Use Template ---
        source_notice = f"This RFP was generated from the '{template_path.name}' template."
        with open(template_path, "r") as f:
            rfp_text = f.read()

        # Replace placeholders
        for key, value in initiative_data.items():
            # Handle list values by joining them
            if isinstance(value, list):
                value_str = ", ".join(map(str, value))
            else:
                value_str = str(value)
            rfp_text = rfp_text.replace(f"{{{{{key}}}}}", value_str)

        # Special placeholder for current date
        rfp_text = rfp_text.replace("{% raw %}{{% endraw %}CURRENT_DATE{% raw %}}{% endraw %}", date.today().isoformat())

    else:
        # --- Fallback to LLM ---
        source_notice = "This RFP was generated by Gemini. Review and edit as needed."
        if not gemini_model:
            return HTMLResponse("<h3>Gemini API is not configured. Please set the GOOGLE_API_KEY environment variable.</h3>", status_code=500)

        prompt = f"""
        Based on the following sourcing initiative data, generate a professional and comprehensive Request for Proposal (RFP) document.
        The document should be well-structured with clear sections, headings, and lists.

        Sourcing Initiative Data:
        {json.dumps(initiative_data, indent=2)}
        """
        response = gemini_model.generate_content(prompt)
        rfp_text = response.text

    # Save docx for download
    save_rfp_doc(rfp_text, initiative_id)
    safe_text = html_lib.escape(rfp_text)
    html = '<div class="container">'
    html += render_progress(3)
    html += "<h1>📄 Generated RFP</h1>"
    html += f'<div class="rfp-output">{safe_text}</div>'
    html += f'<a class="download" href="/download_rfp/{initiative_id}">⬇️ Download as Word (.docx)</a>'
    html += f'<p class="notice">{source_notice}</p>'
    html += f'<a class="download" href="/upload_vendor_responses/{initiative_id}">⬆️ Upload Vendor Responses</a>'
    html += '</div>'
    return get_base_layout(f"RFP for Initiative #{initiative_id}", html)

@app.get("/find_vendors/{initiative_id}/{schema_name}", response_class=HTMLResponse)
async def find_vendors_loading(initiative_id: int, schema_name: str):
    """Show a loading screen while the AI searches for vendors."""
    html = '<div class="container">'
    html += "<h1>🤖 Finding Potential Vendors...</h1>"
    html += '<div class="loader"></div>'
    html += "<p class='notice'>The AI is analyzing your requirements to suggest suitable vendors. This may take a moment.</p>"
    # JS redirect to the result endpoint
    html += f"""
    <script>
      setTimeout(function(){{ window.location.href = '/find_vendors_result/{initiative_id}/{schema_name}'; }}, 800);
    </script>
    """
    html += '</div>'
    return get_base_layout("Finding Vendors...", html)


@app.get("/find_vendors_result/{initiative_id}/{schema_name}", response_class=HTMLResponse)
async def find_vendors_result(initiative_id: int, schema_name: str):
    """Use Ollama to find vendors based on initiative data."""
    try:
        initiative_data = load_initiative_data(initiative_id, schema_name)
    except FileNotFoundError:
        return HTMLResponse("<h3>Initiative data not found.</h3>", status_code=404)

    prompt = f"""
You are a pharmaceutical industry sourcing specialist. Based on the following project details, please identify and list 7 potential vendors that would be a good fit.

For each vendor, provide a brief (1-2 sentence) justification for why they are a good match based on the project requirements.

Project Details:
{json.dumps(initiative_data, indent=2)}

Please format your response as a list.
"""

    if not gemini_model:
        return HTMLResponse("<h3>Gemini API is not configured. Please set the GOOGLE_API_KEY environment variable.</h3>", status_code=500)

    try:
        prompt="What is Capital of Italy ?"
        response = gemini_model.generate_content(prompt) # Use the detailed prompt
        result_text = response.text
    except Exception as e:
        error_message = f"<h3>Error calling Gemini API:</h3><pre>{html_lib.escape(str(e))}</pre>"
        return HTMLResponse(error_message, status_code=500)

    html = '<div class="container">'
    html += "<h1>🤖 Suggested Vendors</h1>"
    html += f'<div class="rfp-output" style="white-space: pre-wrap;">{html_lib.escape(result_text)}</div>'
    html += f'<p class="notice">These vendors were suggested by Gemini based on your input. Further vetting is recommended.</p>'
    html += f'<p><a href="/rfp_result/{initiative_id}/{schema_name}">← Back to RFP</a></p>'
    html += '</div>'
    return get_base_layout(f"Vendors for Initiative #{initiative_id}", html)

@app.get("/download_rfp/{initiative_id}")
async def download_rfp(initiative_id:int):
    path = RFP_FOLDER / f"initiative_{initiative_id}_rfp.docx"
    if path.exists():
        return FileResponse(str(path), filename=f"initiative_{initiative_id}_RFP.docx")
    return HTMLResponse("<h3>RFP not found.</h3>", status_code=404)

@app.get("/initiative/{initiative_id}", response_class=JSONResponse)
async def get_initiative(initiative_id:int):
    # return base submission if exists
    file = SUBMISSION_FOLDER / f"initiative_{initiative_id}.json"
    if not file.exists():
        return JSONResponse({"error":"Initiative not found"}, status_code=404)
    with open(file,"r") as f:
        return JSONResponse(json.load(f))

@app.get("/health")
async def health():
    return {"status":"ok"}



@app.get("/upload_vendor_responses/{initiative_id}", response_class=HTMLResponse)
async def upload_vendor_form(initiative_id: int):
    """Show upload form for vendor responses."""
    html_content = f"""
    <div class="container">
        <h1>📤 Upload Vendor Responses</h1>
        <form action="/upload_vendor_responses/{initiative_id}" method="post" enctype="multipart/form-data">
            <p>Please upload between <b>2 and 7</b> vendor response files (PDF, DOCX, or TXT).</p>
            <input type="file" name="files" multiple required accept=".pdf,.docx,.txt">
            <br><br>
            <button type="submit">Upload & Compare</button>
        </form>
        <p class="notice">Each vendor response will be analyzed and compared using AI.</p>
    </div>
    """
    return get_base_layout(f"Upload Responses for Initiative #{initiative_id}", html_content)


VENDOR_FOLDER = Path("data/vendor_responses")
VENDOR_FOLDER.mkdir(parents=True, exist_ok=True)

@app.post("/upload_vendor_responses/{initiative_id}")
async def upload_vendor_files(initiative_id: int, files: list[UploadFile] = File(...)):
    """Upload and process vendor responses, enforcing 2–7 file count."""
    if len(files) < 2 or len(files) > 7:
        error_html = f"""
            <div class='container'>
                <h2>⚠️ Invalid number of files uploaded</h2>
                <p>You must upload between <b>2 and 7</b> vendor responses. You uploaded {len(files)}.</p>
                <a href="/upload_vendor_responses/{initiative_id}">← Try Again</a>
            </div>
            """
        return get_base_layout("Upload Error", error_html)

    upload_dir = VENDOR_FOLDER / f"initiative_{initiative_id}"
    upload_dir.mkdir(parents=True, exist_ok=True)

    combined_data = {}

    for file in files:
        content = await file.read()
        text = ""

        if file.filename.lower().endswith(".pdf"):
            pdf = PdfReader(io.BytesIO(content))
            for page in pdf.pages:
                text += page.extract_text() or ""
        elif file.filename.lower().endswith(".docx"):
            doc = Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            text = content.decode("utf-8", errors="ignore")

        combined_data[file.filename] = text.strip()
        with open(upload_dir / file.filename, "wb") as f:
            f.write(content)

    combined_path = upload_dir / "combined_vendor_responses.json"
    with open(combined_path, "w") as f:
        json.dump(combined_data, f, indent=2)

    # Call compare page handler to run the AI comparison immediately and return its HTML
    return RedirectResponse(url=f"/compare_vendors/{initiative_id}", status_code=303)

@app.get("/download_comparison/{initiative_id}")
async def download_comparison(initiative_id: int):
    """Download vendor comparison result."""
    result_path = VENDOR_FOLDER / f"initiative_{initiative_id}" / "comparison_result.txt"
    if not result_path.exists():
        return HTMLResponse("<h3>Comparison result not found.</h3>", status_code=404)
    return FileResponse(str(result_path), filename=f"initiative_{initiative_id}_comparison.txt")

@app.get("/download_comparison_docx/{initiative_id}")
async def download_comparison_docx(initiative_id: int):
    """Download vendor comparison result as .docx."""
    result_path = VENDOR_FOLDER / f"initiative_{initiative_id}" / "comparison_result.docx"
    if not result_path.exists():
        return HTMLResponse("<h3>Word comparison result not found.</h3>", status_code=404)
    return FileResponse(str(result_path), filename=f"initiative_{initiative_id}_comparison.docx")

@app.get("/download_comparison_xlsx/{initiative_id}")
async def download_comparison_xlsx(initiative_id: int):
    """Download vendor comparison result as .xlsx."""
    result_path = VENDOR_FOLDER / f"initiative_{initiative_id}" / "comparison_result.xlsx"
    if not result_path.exists():
        return HTMLResponse("<h3>Excel comparison result not found.</h3>", status_code=404)
    return FileResponse(str(result_path), filename=f"initiative_{initiative_id}_comparison.xlsx")

@app.get("/compare_vendors/{initiative_id}", response_class=HTMLResponse)
async def compare_vendors_page(initiative_id: int):
    """Compare vendor responses using Ollama AI."""
    combined_path = VENDOR_FOLDER / f"initiative_{initiative_id}" / "combined_vendor_responses.json"
    if not combined_path.exists():
        return HTMLResponse("<h3>No vendor responses uploaded yet.</h3>", status_code=404)

    try:
        with open(combined_path) as f:
            vendor_data = json.load(f)

        # Construct prompt for AI comparison (string)
        prompt = f"""
You are an expert RFP evaluation specialist. Your task is to analyze and compare the following vendor responses for initiative {initiative_id}.

**Instructions:**
1.  Carefully review each vendor's response text provided in the JSON below.
2.  For each vendor, provide a concise summary, list their key strengths and weaknesses, and identify any potential risks.
3.  Score each vendor on a scale of 0 to 10 for the following criteria:
    - Technical Capability
    - Quality & Compliance
    - Project Management
    - Supply Reliability
    - Cost Competitiveness
4.  Calculate a percentage for each score (score / 10 * 100).
5.  Provide an overall recommendation, including a summary of why you are recommending the top vendors.
6.  Format your entire output as a single, valid JSON object. Do not include any text or formatting outside of the JSON block.

Vendor Responses:
{json.dumps(vendor_data, indent=2)}

**JSON Output Structure:**
```json
{{
  "vendors": [
    {{
      "vendor_name": "Vendor A Name",
      "summary": "A brief summary of Vendor A's proposal.",
      "scores": {{
        "Technical Capability": {{"score": 8, "percentage": 80}},
        "Quality & Compliance": {{"score": 9, "percentage": 90}},
        "Project Management": {{"score": 7, "percentage": 70}},
        "Supply Reliability": {{"score": 8, "percentage": 80}},
        "Cost Competitiveness": {{"score": 6, "percentage": 60}}
      }},
      "strengths": "List of strengths for Vendor A.",
      "weaknesses": "List of weaknesses for Vendor A.",
      "risks": "Identified risks for Vendor A."
    }}
  ],
  "recommendation": {{
    "summary": "Overall summary of the evaluation and justification for the recommendation.",
    "top_vendors": ["Vendor A Name", "Vendor B Name"]
  }}
}}
```
        """
        if not gemini_model:
            return HTMLResponse("<h3>Gemini API is not configured. Please set the GOOGLE_API_KEY environment variable.</h3>", status_code=500)
        response = gemini_model.generate_content(prompt)
        result_text_raw = response.text

        # Save raw text and structured data
        (VENDOR_FOLDER / f"initiative_{initiative_id}" / "comparison_result.txt").write_text(result_text_raw)
        parsed_data = json.loads(result_text_raw)
        save_comparison_docx(parsed_data, initiative_id)
        save_comparison_xlsx(parsed_data, initiative_id)

    except Exception as e:
        error_message = f"<h3>Error calling Gemini API:</h3><pre>{html_lib.escape(str(e))}</pre>"
        return HTMLResponse(error_message, status_code=500)

    result_text_safe = html_lib.escape(json.dumps(parsed_data, indent=2))

    html_content = f"""
    <div class="container">
        <h1>🏁 Vendor Comparison Results</h1>
        <div class="rfp-output">{result_text_safe}</div>
        <a class="download" href="/download_comparison_docx/{initiative_id}">⬇️ Download as Word (.docx)</a>
        <a class="download" href="/download_comparison_xlsx/{initiative_id}">⬇️ Download as Excel (.xlsx)</a>
        <a class="download" href="/download_comparison/{initiative_id}">⬇️ Download Results (.txt)</a>
        <p><a href="/initiatives">← Back to Initiatives</a></p>
    </div>
    """
    return get_base_layout(f"Comparison for Initiative #{initiative_id}", html_content)

#
# ---- end of app.py ----
